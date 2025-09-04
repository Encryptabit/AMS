#!/usr/bin/env python3
"""
Simple script to create a test DOCX file for testing the book parser.
Requires python-docx: pip install python-docx
"""

from docx import Document
from docx.shared import Inches
import os

# Create a new Document
doc = Document()

# Add title
title = doc.add_heading('My Test Book', 0)

# Add some paragraphs
p1 = doc.add_paragraph('This is the first paragraph of our test DOCX file. ')
p1.add_run('It contains multiple sentences to test the DOCX parsing capabilities.')

p2 = doc.add_paragraph('This is the second paragraph with different content. ')
p2.add_run('We want to ensure that the parser can extract text from Word documents correctly.')

# Add a chapter heading
doc.add_heading('Chapter 1: The Beginning', level=1)

p3 = doc.add_paragraph('The story begins here with the first chapter. ')
p3.add_run('This paragraph introduces the main concepts we want to validate.')

p4 = doc.add_paragraph('Finally, we can test how the book indexer will split this DOCX content ')
p4.add_run('into words, sentences, and paragraphs for later validation against audio transcripts.')

# Save the document
output_path = 'test_book.docx'
doc.save(output_path)
print(f"Test DOCX file created: {output_path}")

# Set document properties
doc.core_properties.title = 'My Test Book'
doc.core_properties.author = 'Test Author'
doc.core_properties.subject = 'Book parsing test'
doc.core_properties.keywords = 'test, book, parsing'

# Save again with properties
doc.save(output_path)
print(f"DOCX file saved with metadata: {output_path}")